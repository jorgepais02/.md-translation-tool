from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches


# ---------------------------
# Constants
# ---------------------------

PROJECT_ROOT = Path(__file__).resolve().parent.parent
HEADER_IMAGE = PROJECT_ROOT / "public" / "header.png"

# Languages that use right-to-left script
RTL_LANGS = {"ar", "he", "fa", "ur"}

# Font selection by language
FONT_MAP = {
    "zh": "SimSun",          # Chinese – common CJK serif
    "ar": "Traditional Arabic",  # Arabic – elegant Naskh
}
DEFAULT_FONT = "Times New Roman"


# ---------------------------
# Layout configuration
# ---------------------------

def _font_for(lang: str) -> str:
    return FONT_MAP.get(lang, DEFAULT_FONT)


def _is_rtl(lang: str) -> bool:
    return lang in RTL_LANGS


def set_margins_standard(doc: Document) -> None:
    """Standard academic margins: 2.5 cm."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_style_base(doc: Document, lang: str = "es") -> None:
    """Configure base typography. Adapts font for language."""
    font_name = _font_for(lang)

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(12)
    normal.font.bold = False
    normal.font.color.rgb = RGBColor(0, 0, 0)

    h1 = doc.styles["Heading 1"]
    h1.font.name = font_name
    h1.font.size = Pt(20)
    h1.font.bold = False
    h1.font.color.rgb = RGBColor(0, 0, 0)

    h2 = doc.styles["Heading 2"]
    h2.font.name = font_name
    h2.font.size = Pt(16)
    h2.font.bold = False
    h2.font.color.rgb = RGBColor(0, 0, 0)


# ---------------------------
# RTL helpers
# ---------------------------

def _set_paragraph_rtl(p) -> None:
    """Mark a paragraph as right-to-left at the XML level."""
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _set_run_rtl(run) -> None:
    """Mark a run as right-to-left and set complex-script font."""
    rPr = run._r.get_or_add_rPr()
    rtl_elem = OxmlElement("w:rtl")
    rtl_elem.set(qn("w:val"), "1")
    rPr.append(rtl_elem)


def _apply_rtl(p, lang: str) -> None:
    """Apply RTL paragraph direction + run flags if language requires it."""
    if not _is_rtl(lang):
        return
    _set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        _set_run_rtl(run)


# ---------------------------
# Paragraph formatting rules
# ---------------------------

BODY_LINE_SPACING = 1.2
BODY_SPACE_AFTER_PT = 10
TITLE_SPACE_AFTER_PT = 36

H1_SPACE_BEFORE_PT = 28
H1_SPACE_AFTER_PT = 6
H2_SPACE_BEFORE_PT = 20
H2_SPACE_AFTER_PT = 5

SPACE_AFTER_BEFORE_HEADING_PT = 2

LIST_ITEM_SPACE_BEFORE_PT = 0
LIST_ITEM_SPACE_AFTER_PT = 6
LIST_BLOCK_END_SPACE_AFTER_PT = 12

LIST_BASE_LEFT_INDENT_CM = 1.5
LIST_LEVEL_INDENT_CM = 0.5
LIST_HANGING_CM = 0.4

BULLET_INDENT_SPACES = 2


def format_body(p, lang: str = "es") -> None:
    font_name = _font_for(lang)
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = BODY_LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(BODY_SPACE_AFTER_PT)

    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(12)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)

    _apply_rtl(p, lang)


def format_heading(p, level: int, lang: str = "es") -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if level == 1:
        p.paragraph_format.space_before = Pt(H1_SPACE_BEFORE_PT)
        p.paragraph_format.space_after = Pt(H1_SPACE_AFTER_PT)
    elif level == 2:
        p.paragraph_format.space_before = Pt(H2_SPACE_BEFORE_PT)
        p.paragraph_format.space_after = Pt(H2_SPACE_AFTER_PT)
    else:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)

    for run in p.runs:
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)

    _apply_rtl(p, lang)


def trim_space_after_last_body_paragraph(doc: Document, pt: int = SPACE_AFTER_BEFORE_HEADING_PT) -> None:
    if not doc.paragraphs:
        return
    last = doc.paragraphs[-1]
    style_name = getattr(getattr(last, "style", None), "name", "") or ""
    if style_name == "Normal":
        last.paragraph_format.space_after = Pt(pt)


def trim_space_before_list(doc: Document, pt: int = 0) -> None:
    if not doc.paragraphs:
        return
    last = doc.paragraphs[-1]
    style_name = getattr(getattr(last, "style", None), "name", "") or ""
    if style_name == "Normal":
        last.paragraph_format.space_after = Pt(pt)


def add_space_after_list_if_needed(doc: Document) -> None:
    if not doc.paragraphs:
        return
    last = doc.paragraphs[-1]
    style_name = getattr(getattr(last, "style", None), "name", "") or ""
    if "List" in style_name:
        last.paragraph_format.space_after = Pt(LIST_BLOCK_END_SPACE_AFTER_PT)


# ---------------------------
# Header image
# ---------------------------

def add_header_image(doc: Document) -> None:
    """Insert header.png spanning the full page width (edge to edge)."""
    if not HEADER_IMAGE.exists():
        return

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remove paragraph spacing so the image sits flush
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Negative indents to compensate for page margins → true edge-to-edge
    left_margin = section.left_margin
    right_margin = section.right_margin
    p.paragraph_format.left_indent = -left_margin
    p.paragraph_format.right_indent = -right_margin

    # Clear existing content
    for run in p.runs:
        run.clear()

    run = p.add_run()
    # Full A4 page width: 21 cm
    page_width = section.page_width
    run.add_picture(str(HEADER_IMAGE), width=page_width)


# ---------------------------
# Footer page number
# ---------------------------

def add_page_number_footer_right(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)

    run.font.name = DEFAULT_FONT
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)


# ---------------------------
# Title / Headings / Body
# ---------------------------

def add_title(doc: Document, text: str, lang: str = "es") -> None:
    font_name = _font_for(lang)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text.strip())
    run.bold = False
    run.font.name = font_name
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)

    p.paragraph_format.space_after = Pt(TITLE_SPACE_AFTER_PT)
    _apply_rtl(p, lang)


def add_heading(doc: Document, text: str, md_level: int, lang: str = "es") -> None:
    word_level = max(1, md_level - 1)
    word_level = min(word_level, 9)
    p = doc.add_heading(text.strip(), level=word_level)
    format_heading(p, level=word_level, lang=lang)


def add_body(doc: Document, text: str, lang: str = "es") -> None:
    p = doc.add_paragraph(text.rstrip())
    format_body(p, lang=lang)


# ---------------------------
# Lists with "bold before ':'"
# ---------------------------

def _add_runs_bold_before_colon(p, text: str, lang: str = "es") -> None:
    font_name = _font_for(lang)
    if ":" in text:
        first, rest = text.split(":", 1)

        run_bold = p.add_run(first.strip() + ":")
        run_bold.bold = True
        run_bold.font.name = font_name
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)

        run_normal = p.add_run(" " + rest.strip())
        run_normal.bold = False
        run_normal.font.name = font_name
        run_normal.font.size = Pt(12)
        run_normal.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(text.strip())
        run.bold = False
        run.font.name = font_name
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)


def _apply_list_indents(p, level: int, lang: str = "es") -> None:
    if _is_rtl(lang):
        # RTL: indent from the right side
        right = LIST_BASE_LEFT_INDENT_CM + (level * LIST_LEVEL_INDENT_CM)
        p.paragraph_format.right_indent = Cm(right)
        p.paragraph_format.first_line_indent = Cm(-LIST_HANGING_CM)
    else:
        left = LIST_BASE_LEFT_INDENT_CM + (level * LIST_LEVEL_INDENT_CM)
        p.paragraph_format.left_indent = Cm(left)
        p.paragraph_format.first_line_indent = Cm(-LIST_HANGING_CM)


def add_bullet(doc: Document, text: str, level: int = 0, lang: str = "es") -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2" if level == 1 else "List Bullet 3"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = BODY_LINE_SPACING
    p.paragraph_format.space_before = Pt(LIST_ITEM_SPACE_BEFORE_PT)
    p.paragraph_format.space_after = Pt(LIST_ITEM_SPACE_AFTER_PT)

    _apply_list_indents(p, level, lang=lang)
    _add_runs_bold_before_colon(p, text, lang=lang)
    _apply_rtl(p, lang)


def add_numbered(doc: Document, text: str, level: int = 0, lang: str = "es") -> None:
    style = "List Number" if level == 0 else "List Number 2" if level == 1 else "List Number 3"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = BODY_LINE_SPACING
    p.paragraph_format.space_before = Pt(LIST_ITEM_SPACE_BEFORE_PT)
    p.paragraph_format.space_after = Pt(LIST_ITEM_SPACE_AFTER_PT)

    _apply_list_indents(p, level, lang=lang)
    _add_runs_bold_before_colon(p, text, lang=lang)
    _apply_rtl(p, lang)


# ---------------------------
# Markdown parsing
# ---------------------------

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")


def md_to_docx(md_path: Path, docx_path: Path, lang: str = "es") -> None:
    md_lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    set_margins_standard(doc)
    set_style_base(doc, lang=lang)
    add_header_image(doc)
    add_page_number_footer_right(doc)

    # Set document-level RTL if needed
    if _is_rtl(lang):
        for section in doc.sections:
            sectPr = section._sectPr
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            sectPr.append(bidi)

    title_written = False
    buffer: list[str] = []

    def flush_buffer() -> None:
        nonlocal buffer
        para = " ".join(line.strip() for line in buffer).strip()
        buffer = []
        if para:
            add_space_after_list_if_needed(doc)
            add_body(doc, para, lang=lang)

    for line in md_lines:
        m = HEADING_RE.match(line)
        if m:
            flush_buffer()
            hashes, heading_text = m.group(1), m.group(2)
            level = len(hashes)

            if level == 1 and not title_written:
                add_title(doc, heading_text, lang=lang)
                title_written = True
            else:
                trim_space_after_last_body_paragraph(doc)
                add_heading(doc, heading_text, md_level=level, lang=lang)
            continue

        if line.strip() == "":
            flush_buffer()
            continue

        # LIST DETECTION
        raw = line.rstrip("\n")
        leading = len(raw) - len(raw.lstrip(" \t"))
        leading_spaces = raw[:leading].replace("\t", " " * 4)
        indent = len(leading_spaces)
        level = indent // BULLET_INDENT_SPACES
        level = min(level, 2)

        s = raw.lstrip(" \t").strip()

        # Numbered list
        m_num = re.match(r"^(\d+)\.\s+(.*\S)\s*$", s)
        if m_num:
            flush_buffer()
            trim_space_before_list(doc, pt=0)
            content = m_num.group(2).strip()
            add_numbered(doc, content, level=level, lang=lang)
            continue

        # Bullet list
        if s.startswith("- "):
            flush_buffer()
            trim_space_before_list(doc, pt=0)
            content = s[2:].strip()

            if " - " in content:
                parts = [p.strip() for p in content.split(" - ") if p.strip()]
                for part in parts:
                    if part.startswith("- "):
                        part = part[2:].strip()
                    add_bullet(doc, part, level=level, lang=lang)
            else:
                add_bullet(doc, content, level=level, lang=lang)
            continue

        buffer.append(line)

    flush_buffer()
    doc.save(docx_path)


# ---------------------------
# CLI
# ---------------------------

def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Uso:")
        print("  python make_notes.py apuntes.md [salida.docx] [--lang es]")
        return 2

    md_path = Path(argv[1]).expanduser().resolve()
    if not md_path.exists():
        print(f"ERROR: No existe {md_path}")
        return 2

    out = Path(argv[2]).expanduser().resolve() if len(argv) >= 3 else md_path.with_suffix(".docx")

    # Parse optional --lang flag
    lang = "es"
    if "--lang" in argv:
        idx = argv.index("--lang")
        if idx + 1 < len(argv):
            lang = argv[idx + 1].lower()

    md_to_docx(md_path, out, lang=lang)

    print(f"OK -> {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
